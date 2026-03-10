from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse, JsonResponse
from django.contrib.auth.models import User
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from reportlab.pdfgen import canvas
from docx import Document
from .models import Task, TaskFile, UserProfile   #  Removed UploadedCSV import

# Home page view
def home(request):
    # Simple welcome message (can be replaced with home.html template)
    return HttpResponse("<h1>Welcome to PDMS</h1><p>Go to <a href='/pdms/downloads/'>Downloads</a></p>")

# About page view
def about_page(request):
    # Render about.html template
    return render(request, "pdms_app/about.html")

# Downloads page view
def download_page(request):
    # Render download.html template with upload + export buttons
    return render(request, "pdms_app/download.html")

# Task detail view (shows task info + attached files)
def task_detail(request, task_id):
    task = get_object_or_404(Task, id=task_id)  # Get task or 404 if not found
    return render(request, "pdms_app/task_detail.html", {"task": task})

# File upload view for a specific task
def upload_task_file(request, task_id):
    task = get_object_or_404(Task, id=task_id)  # Get task
    if request.method == "POST" and request.FILES.get("file"):  # Check if file uploaded
        TaskFile.objects.create(task=task, file=request.FILES["file"])  # Save file
        return redirect("task_detail", task_id=task.id)  # Redirect back to task detail
    return HttpResponse("No file uploaded", status=400)

# User dashboard view (shows profile, teams, assigned tasks)
def user_dashboard(request, user_id):
    user = get_object_or_404(User, id=user_id)  # Get user
    profile = get_object_or_404(UserProfile, user=user)  # Get user profile
    tasks = Task.objects.filter(assigned_to=user)  # Get tasks assigned to user

    context = {
        "user": user,
        "profile": profile,
        "tasks": tasks,
    }
    return render(request, "pdms_app/dashboard.html", context)

# API endpoint: upload tasks via CSV
@api_view(['POST'])
@permission_classes([IsAuthenticated])  # Require JWT authentication
def upload_csv(request):
    if request.FILES.get("file"):
        uploaded_file = request.FILES["file"]

        # Process file directly without saving to DB
        decoded_file = uploaded_file.read().decode("utf-8").splitlines()
        import csv
        reader = csv.DictReader(decoded_file)
        rows = [row for row in reader]

        return JsonResponse({"message": "CSV uploaded successfully", "rows": rows})
    return JsonResponse({"error": "No file uploaded"}, status=400)

# API endpoint: export tasks as CSV
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def export_csv(request):
    response = HttpResponse(content_type="text/csv")
    response["Content-Disposition"] = 'attachment; filename="pdms_tasks.csv"'
    import csv
    writer = csv.writer(response)
    writer.writerow(["Task ID", "Title", "Description"])
    for task in Task.objects.all():
        writer.writerow([task.id, task.title, task.description])
    return response

# API endpoint: export tasks as PDF
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def export_pdf(request):
    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="pdms_tasks.pdf"'
    p = canvas.Canvas(response)
    p.drawString(100, 800, "PDMS Task Report")
    y = 750
    for task in Task.objects.all():
        p.drawString(100, y, f"{task.id} - {task.title}: {task.description}")
        y -= 20
    p.showPage()
    p.save()
    return response

# API endpoint: export tasks as DOCX
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def export_docx(request):
    document = Document()
    document.add_heading("PDMS Task Report", 0)
    for task in Task.objects.all():
        document.add_paragraph(f"{task.id} - {task.title}: {task.description}")
    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = 'attachment; filename="pdms_tasks.docx"'
    document.save(response)
    return response
